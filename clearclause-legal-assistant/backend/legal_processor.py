import os
import time
from typing import List
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
from docx import Document
import fitz  # PyMuPDF
from dotenv import load_dotenv
from pathlib import Path
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Load .env from this backend directory explicitly (and override existing)
_env_path = Path(__file__).resolve().parent / ".env"
load_dotenv(dotenv_path=_env_path, override=True)

# Configure Gemini API
# Accept either GEMINI_API_KEY or GOOGLE_API_KEY and propagate to both for library compatibility
api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
if api_key:
    # Set both env vars so downstream libs (langchain-google-genai) pick it up
    os.environ["GOOGLE_API_KEY"] = api_key
    os.environ["GEMINI_API_KEY"] = api_key
    genai.configure(api_key=api_key)
else:
    print("Warning: Google API Key not found. Please set GEMINI_API_KEY or GOOGLE_API_KEY in backend/.env")

# --- Helper Functions ---

def safe_extract_text(response):
    """Safely extract text from Gemini response"""
    try:
        if hasattr(response, "text") and response.text:
            return response.text
        if hasattr(response, "candidates") and response.candidates:
            parts = response.candidates[0].content.parts
            if parts and hasattr(parts[0], "text"):
                return parts[0].text
    except Exception:
        pass
    return "(No content returned)"

def robust_generate(prompt, retries=2):
    """Robust text generation with retry logic"""
    models_to_try = ["gemini-2.0-flash-exp", "gemini-1.5-flash"]
    
    for model_name in models_to_try:
        try:
            model = genai.GenerativeModel(model_name)
            for attempt in range(retries + 1):
                try:
                    resp = model.generate_content(prompt)
                    text = safe_extract_text(resp)
                    if text.strip() and "No content returned" not in text:
                        return text
                except Exception as e:
                    if "429" in str(e):  # Quota exceeded
                        break
                    if attempt < retries:
                        time.sleep(1)
                        continue
                    return f"(Error: {e})"
        except:
            continue
    
    return "(No model could return a valid response)"

def translate_text(text, target_language):
    """Translate text to target language with offline fallback (Hindi/Kannada)."""
    prompt = (
        f"Translate the following English legal/business text to {target_language}. "
        "Provide only the translation without any explanations or extra commentary.\n\n" + text
    )
    result = robust_generate(prompt)
    # If API quota or error produced an error-like message, use offline fallback
    lowered = result.lower()
    if "quota" in lowered or "error" in lowered or "no model" in lowered:
        return offline_translate(text, target_language)
    return result

def offline_translate(text: str, target_language: str) -> str:
    """Very lightweight dictionary-based fallback translation for key legal terms."""
    # Extended term dictionaries; phrases processed first.
    hindi_terms = {
        # Multi-word phrases
        "governing law": "प्रभावी कानून",
        "this agreement": "यह समझौता",
        "mutual consent": "पारस्परिक सहमति",
        "shall survive": "जारी रहेगा",
        "force majeure": "दैविक आपदा",
        "in witness whereof": "इस साक्ष्य में",
        "non disclosure": "गैर प्रकटीकरण",
        "data protection": "डेटा संरक्षण",
        "limitation of liability": "देयता की सीमा",
        "applicable law": "प्रासंगिक कानून",
        # Single words
        "agreement": "समझौता",
        "party": "पक्ष",
        "parties": "पक्षों",
        "confidentiality": "गोपनीयता",
        "obligation": "दायित्व",
        "obligations": "दायित्वों",
        "termination": "समापन",
        "year": "वर्ष",
        "years": "वर्ष",
        "governing": "प्रभावी",
        "law": "कानून",
        "india": "भारत",
        "signed": "हस्ताक्षरित",
        "electronic": "इलेक्ट्रॉनिक",
        "document": "दस्तावेज",
        "clause": "धारा",
        "clauses": "धाराएँ",
        "data": "डेटा",
        "privacy": "गोपनीयता",
        "liability": "देयता",
        "shall": "करेंगे",
        "may": "सकते हैं",
        "must": "अवश्य",
        "including": "समेत",
        "disclosure": "प्रकटीकरण",
        "effective": "प्रभावी",
        "date": "तिथि",
        "warranty": "वारंटी",
        "indemnity": "प्रतिपूर्ति",
    }
    kannada_terms = {
        # Multi-word phrases
        "governing law": "ಪ್ರಭಾವಿ ಕಾನೂನು",
        "this agreement": "ಈ ಒಪ್ಪಂದ",
        "mutual consent": "ಪರಸ್ಪರ ಒಪ್ಪಿಗೆ",
        "shall survive": "ಮುಂದುವರಿಯುತ್ತದೆ",
        "force majeure": "ಅಪರಿಹಾರ್ಯ ಪರಿಸ್ಥಿತಿ",
        "in witness whereof": "ಸಾಕ್ಷಿಯಾಗಿ",
        "non disclosure": "ಅಪ್ರಕಟಣೆ",
        "data protection": "ಡೇಟಾ ರಕ್ಷಣಾ",
        "limitation of liability": "ಬಾಧ್ಯತೆಯ ಮಿತಿ",
        "applicable law": "ಅನ್ವಯಿಸುವ ಕಾನೂನು",
        # Single words
        "agreement": "ಒಪ್ಪಂದ",
        "party": "ಪಕ್ಷ",
        "parties": "ಪಕ್ಷಗಳು",
        "confidentiality": "ಗೌಪ್ಯತೆ",
        "obligation": "ಬದ್ದತೆ",
        "obligations": "ಬದ್ದತೆಗಳು",
        "termination": "ರದ್ದು",
        "year": "ವರ್ಷ",
        "years": "ವರ್ಷಗಳು",
        "governing": "ಪ್ರಭಾವಿ",
        "law": "ಕಾನೂನು",
        "india": "ಭಾರತ",
        "signed": "ಸಹಿ ಮಾಡಲಾಗಿದೆ",
        "electronic": "ಇಲೆಕ್ಟ್ರಾನಿಕ್",
        "document": "ಡಾಕ್ಯುಮೆಂಟ್",
        "clause": "ವಿಧಾನ",
        "clauses": "ವಿಧಾನಗಳು",
        "data": "ಡೇಟಾ",
        "privacy": "ಗೌಪ್ಯತೆ",
        "liability": "ಬಾಧ್ಯತೆ",
        "shall": "ಮಾಡಬೇಕು",
        "may": "ಬಹುದಾಗಿದೆ",
        "must": "ಅಗತ್ಯ",
        "including": "ಒಳಗೊಂಡಂತೆ",
        "disclosure": "ಪ್ರಕಟಣೆ",
        "effective": "ಪ್ರಭಾವಿ",
        "date": "ದಿನಾಂಕ",
        "warranty": "ಹಾಮಿ",
        "indemnity": "ನಷ್ಟಪೂರೈಕೆ",
    }

    mapping = hindi_terms if target_language == "Hindi" else kannada_terms

    # Phrase replacement (longest first)
    processed = text
    phrase_keys = sorted([k for k in mapping if " " in k], key=len, reverse=True)
    for pk in phrase_keys:
        pattern = re.compile(rf"\b{re.escape(pk)}\b", re.IGNORECASE)
        processed = pattern.sub(mapping[pk], processed)

    # Token-level replacement & transliteration fallback
    tokens = re.findall(r"[A-Za-z]{2,}|\d+|[^A-Za-z\d]", processed)

    def transliterate(word: str) -> str:
        # Enhanced syllable-aware transliteration with consonant+vowel patterns
        w_lower = word.lower()
        
        # Consonant and vowel mappings for Hindi (Devanagari)
        dev_cons = {
            'k':'क','kh':'ख','g':'ग','gh':'घ','ch':'च','chh':'छ','j':'ज','jh':'झ',
            't':'त','th':'थ','d':'द','dh':'ध','n':'न','p':'प','ph':'फ','b':'ब','bh':'भ',
            'm':'म','y':'य','r':'र','l':'ल','v':'व','w':'व','sh':'श','s':'स','h':'ह',
            'c':'क','f':'फ','q':'क','x':'क्स','z':'ज'
        }
        dev_vowel_marks = {'a':'','aa':'ा','i':'ि','ii':'ी','u':'ु','uu':'ू','e':'े','ai':'ै','o':'ो','au':'ौ'}
        dev_vowels = {'a':'अ','aa':'आ','i':'इ','ii':'ई','u':'उ','uu':'ऊ','e':'ए','ai':'ऐ','o':'ओ','au':'औ'}
        
        # Kannada mappings
        kn_cons = {
            'k':'ಕ','kh':'ಖ','g':'ಗ','gh':'ಘ','ch':'ಚ','chh':'ಛ','j':'ಜ','jh':'ಝ',
            't':'ತ','th':'ಥ','d':'ದ','dh':'ಧ','n':'ನ','p':'ಪ','ph':'ಫ','b':'ಬ','bh':'ಭ',
            'm':'ಮ','y':'ಯ','r':'ರ','l':'ಲ','v':'ವ','w':'ವ','sh':'ಶ','s':'ಸ','h':'ಹ',
            'c':'ಕ','f':'ಫ','q':'ಕ','x':'ಕ್ಸ್','z':'ಜ'
        }
        kn_vowel_marks = {'a':'','aa':'ಾ','i':'ಿ','ii':'ೀ','u':'ು','uu':'ೂ','e':'ೆ','ai':'ೈ','o':'ೊ','au':'ೌ'}
        kn_vowels = {'a':'ಅ','aa':'ಆ','i':'ಇ','ii':'ಈ','u':'ಉ','uu':'ಊ','e':'ಎ','ai':'ಐ','o':'ಒ','au':'ಔ'}
        
        cons = dev_cons if target_language == "Hindi" else kn_cons
        marks = dev_vowel_marks if target_language == "Hindi" else kn_vowel_marks
        vowels = dev_vowels if target_language == "Hindi" else kn_vowels
        
        out = []
        i = 0
        while i < len(w_lower):
            # Try 3-char consonant (chh)
            if i+2 < len(w_lower):
                c3 = w_lower[i:i+3]
                if c3 in cons:
                    # Check for following vowel
                    if i+3 < len(w_lower):
                        v2 = w_lower[i+3:i+5] if i+4 < len(w_lower) else ""
                        v1 = w_lower[i+3]
                        if v2 in marks:
                            out.append(cons[c3] + marks[v2])
                            i += 5
                            continue
                        elif v1 in marks:
                            out.append(cons[c3] + marks[v1])
                            i += 4
                            continue
                    out.append(cons[c3])
                    i += 3
                    continue
            
            # Try 2-char consonant (kh, sh, etc)
            if i+1 < len(w_lower):
                c2 = w_lower[i:i+2]
                if c2 in cons:
                    if i+2 < len(w_lower):
                        v2 = w_lower[i+2:i+4] if i+3 < len(w_lower) else ""
                        v1 = w_lower[i+2]
                        if v2 in marks:
                            out.append(cons[c2] + marks[v2])
                            i += 4
                            continue
                        elif v1 in marks:
                            out.append(cons[c2] + marks[v1])
                            i += 3
                            continue
                    out.append(cons[c2])
                    i += 2
                    continue
            
            # Try 1-char consonant
            c1 = w_lower[i]
            if c1 in cons:
                if i+1 < len(w_lower):
                    v2 = w_lower[i+1:i+3] if i+2 < len(w_lower) else ""
                    v1 = w_lower[i+1]
                    if v2 in marks:
                        out.append(cons[c1] + marks[v2])
                        i += 3
                        continue
                    elif v1 in marks:
                        out.append(cons[c1] + marks[v1])
                        i += 2
                        continue
                out.append(cons[c1])
                i += 1
                continue
            
            # Standalone vowel
            v2 = w_lower[i:i+2] if i+1 < len(w_lower) else ""
            if v2 in vowels:
                out.append(vowels[v2])
                i += 2
                continue
            if c1 in vowels:
                out.append(vowels[c1])
                i += 1
                continue
            
            # Keep as-is
            out.append(c1)
            i += 1
        
        return "".join(out)

    out_parts = []
    for t in tokens:
        if re.fullmatch(r"[A-Za-z]{2,}", t):
            lower = t.lower()
            out_parts.append(mapping.get(lower, transliterate(lower)))
        else:
            out_parts.append(t)

    out_text = "".join(out_parts)
    # Simple cleanup: collapse spaces
    out_text = re.sub(r"\s+", " ", out_text).strip()
    return out_text + "\n(Offline translation; some words approximated.)"

def _split_sentences(text: str) -> List[str]:
    # Simple sentence splitter; good enough for fallback
    parts = re.split(r"(?<=[.!?])\s+", text.strip())
    return [p.strip() for p in parts if p.strip()]

def _word_freqs(text: str) -> dict:
    words = re.findall(r"[a-zA-Z0-9']+", text.lower())
    stop = set("""
    a an the and or but if then than so because while of in on at to for with from by is are was were be been being this that those these it its it's as not no nor do does did done can could would should shall will may might must also into over under about across between within without against per each other more most some any few many such own same just own very via etc
    """.split())
    freqs = {}
    for w in words:
        if w in stop or len(w) < 3:
            continue
        freqs[w] = freqs.get(w, 0) + 1
    return freqs

def extractive_summary(text: str, max_sentences: int = 7) -> str:
    """Cheap extractive summary: rank sentences by word-frequency score."""
    if not text:
        return "(No content to summarize)"
    sents = _split_sentences(text)
    if len(sents) <= max_sentences:
        return " ".join(sents)
    freqs = _word_freqs(text)
    scored = []
    for i, s in enumerate(sents):
        words = re.findall(r"[a-zA-Z0-9']+", s.lower())
        score = sum(freqs.get(w, 0) for w in words)
        # normalize lightly by sentence length to avoid overlong bias
        score = score / (1 + len(words))
        scored.append((score, i, s))
    scored.sort(reverse=True)
    top = sorted(scored[:max_sentences], key=lambda x: x[1])  # keep original order
    summary = " ".join(s for _,__,s in top)
    return summary

# --- Document Processing ---

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"Error extracting PDF: {e}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        document = Document(file_path)
        for para in document.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
    return text

def get_document_text(file_paths: List[str]) -> str:
    """Extract text from multiple documents"""
    combined_text = ""
    for file_path in file_paths:
        if file_path.endswith('.pdf'):
            combined_text += extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            combined_text += extract_text_from_docx(file_path)
    return combined_text

def get_text_chunks(text: str) -> List[str]:
    """Split text into chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=10000,
        chunk_overlap=1000
    )
    return text_splitter.split_text(text)

# --- Main Processor Class ---

class LegalDocumentProcessor:
    def __init__(self):
        self.vector_store = None
        self.raw_text = ""
        self.api_key = api_key
        self.text_chunks: List[str] = []
        self.tfidf_vectorizer = None
        self.tfidf_matrix = None
    
    def process_documents(self, file_paths: List[str]) -> bool:
        """Process uploaded documents and create vector store"""
        try:
            # Extract text
            raw_text = get_document_text(file_paths)
            if not raw_text.strip():
                print("No text extracted from documents")
                return False
            
            # Create chunks
            text_chunks = get_text_chunks(raw_text)
            if not text_chunks:
                print("Could not create text chunks")
                return False
            
            # Save chunks and text regardless of vector store availability
            self.text_chunks = text_chunks
            self.raw_text = raw_text

            # Build TF-IDF index for offline retrieval
            try:
                self.tfidf_vectorizer = TfidfVectorizer(max_features=500, stop_words='english')
                self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(text_chunks)
            except Exception as e:
                print(f"Could not build TF-IDF index: {e}")
                self.tfidf_vectorizer = None
                self.tfidf_matrix = None

            # Try to build embeddings-backed vector store; if it fails (e.g., quota), fall back to keyword retrieval
            try:
                embeddings = GoogleGenerativeAIEmbeddings(
                    model="models/embedding-001",
                    google_api_key=self.api_key
                )
                self.vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
            except Exception as e:
                # Log and continue with fallback (no vector store)
                print(f"Embedding store unavailable, using fallback retrieval: {e}")
                self.vector_store = None
            
            print(f"Successfully processed {len(file_paths)} document(s)")
            return True
        
        except Exception as e:
            print(f"Error processing documents: {e}")
            return False
    
    def handle_document_qna(self, question: str) -> str:
        """Answer questions about uploaded documents"""
        # Use FAISS when available; else fall back to simple keyword-based retrieval
        if not (self.vector_store or self.text_chunks):
            return "No documents uploaded. Please upload documents first."
        
        try:
            # Search for relevant documents
            if self.vector_store:
                docs = self.vector_store.similarity_search(question, k=3)
            else:
                # Fallback: TF-IDF based retrieval (better than simple keyword matching)
                if self.tfidf_vectorizer and self.tfidf_matrix is not None:
                    q_vec = self.tfidf_vectorizer.transform([question])
                    sims = cosine_similarity(q_vec, self.tfidf_matrix).flatten()
                    top_idxs = np.argsort(sims)[-3:][::-1].tolist()
                else:
                    # Ultimate fallback: keyword scoring
                    q_words = [w for w in question.lower().split() if len(w) > 2]
                    scores = []
                    for idx, ch in enumerate(self.text_chunks):
                        lch = ch.lower()
                        score = sum(lch.count(w) for w in q_words)
                        scores.append((score, idx))
                    scores.sort(reverse=True)
                    top_idxs = [i for _, i in scores[:3] if _ > 0]
                    if not top_idxs:
                        top_idxs = list(range(min(3, len(self.text_chunks))))
                from langchain.schema import Document as LCDocument
                docs = [LCDocument(page_content=self.text_chunks[i]) for i in top_idxs]
            
            # Create QA chain
            prompt_template = """
You are ClearClause, a helpful Legal AI Assistant specialized in analyzing legal documents.
Based on the context provided, answer the question clearly and professionally.

Context:
{context}

Question: {question}

Answer:"""
            
            model = ChatGoogleGenerativeAI(
                model="gemini-2.0-flash-exp",
                temperature=0.3,
                google_api_key=self.api_key
            )
            prompt = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"]
            )
            chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

            # Get answer via model; if that fails due to quota, return relevant excerpts
            try:
                response = chain(
                    {"input_documents": docs, "question": question},
                    return_only_outputs=True
                )
                return response["output_text"]
            except Exception:
                context = "\n\n".join(d.page_content for d in docs)
                # Trim very long context
                if len(context) > 2000:
                    context = context[:2000] + "..."
                return (
                    "(Using offline fallback due to API quota)\n\n"
                    "Relevant excerpts from your document:\n\n" + context
                )
        
        except Exception as e:
            return f"Error answering question: {str(e)}"
    
    def handle_general_qna(self, question: str) -> str:
        """Answer general legal questions"""
        prompt = f"""You are ClearClause, a helpful Legal AI Assistant.
Answer the following legal question clearly and professionally:

Question: {question}

Answer:"""
        return robust_generate(prompt)
    
    def generate_summary(self, custom_instruction: str) -> str:
        """Generate custom summary of documents"""
        if not self.raw_text:
            return "No documents uploaded to summarize."
        
        try:
            # Split into chunks for summarization
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=10000,
                chunk_overlap=1000
            )
            docs = text_splitter.create_documents([self.raw_text])
            
            # Create summarization chain
            model = ChatGoogleGenerativeAI(
                model="gemini-2.0-flash-exp",
                temperature=0.2,
                google_api_key=self.api_key
            )
            chain = load_summarize_chain(model, chain_type="map_reduce")
            
            # Get initial summary
            initial_summary = chain.run(docs)
            
            # Refine with custom instruction
            refine_prompt = f"""
Given the following summary, refine it to follow this instruction: '{custom_instruction}'.

SUMMARY:
{initial_summary}

REFINED SUMMARY:"""
            
            final_summary = robust_generate(refine_prompt)
            return final_summary
        
        except Exception:
            # Fallback: local extractive summary without external API calls
            # Heuristic: longer summary if instruction hints at detailed/point-wise output
            instr = (custom_instruction or "").lower()
            max_sents = 10 if any(k in instr for k in ["detailed", "comprehensive", "bullet", "points"]) else 6
            return extractive_summary(self.raw_text, max_sentences=max_sents)
    
    def clear(self):
        """Clear all data"""
        self.vector_store = None
        self.raw_text = ""
        self.text_chunks = []
        self.tfidf_vectorizer = None
        self.tfidf_matrix = None
