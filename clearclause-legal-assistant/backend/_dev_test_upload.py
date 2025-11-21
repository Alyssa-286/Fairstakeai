from docx import Document
from datetime import datetime
import requests

# Create a small DOCX
path = 'test_doc_upload.docx'
doc = Document()
doc.add_heading('Sample Legal Agreement', level=1)
doc.add_paragraph('This Agreement is made between Party A and Party B on this day.')
doc.add_paragraph('Confidentiality obligations shall survive termination for 2 years.')
doc.add_paragraph('Governing law shall be the laws of India.')
doc.add_paragraph('Signed on {}'.format(datetime.utcnow().date()))
doc.save(path)
print('Created', path)

# Upload the file
url = 'http://localhost:8001/api/upload'
with open(path, 'rb') as f:
    files = {'files': (path, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    r = requests.post(url, files=files)
    print('Status:', r.status_code)
    print('Body:', r.text[:500])
